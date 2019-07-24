### coding: UTF-8
import docx
import datetime
import configparser

# 各定数の取得
const = configparser.ConfigParser()
const.read('conf/const.ini')
const_path = const.get('common','const_path')
file_name = const.get('common','file_name')
old_date = const.get('common','old_date')

now_date = datetime.datetime.today().strftime("%Y%m%d%H%M%S")

# パス
text_path = const_path + file_name + '_' + old_date + '_' + '.txt'
word_path = const_path + file_name + '.docx'
word_bk_path = const_path + 'bk/' + file_name + '_' + now_date + '_' + '.docx'
doc = docx.Document(word_path)
# バックアップ
doc.save(word_bk_path)

try:
    with open(text_path, mode='r') as f:
        word_list = doc.paragraphs
        text_list = f.readlines()
        text_list = [str.rstrip('\n') for str in text_list]
        if len(word_list) < len(text_list) and text_list[len(word_list)-1] == word_list[len(word_list)-1].text:
            for t_index in range(len(word_list), len(text_list)):
                doc.add_paragraph(text_list[t_index])
            doc.save(word_path)
        else:
            raise Exception('想定される処理ではありません。')
except Exception as ex:
    print(ex)
    print('処理を行いませんでした。')
